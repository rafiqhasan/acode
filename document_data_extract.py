from json import load
from docx import Document
from glob import glob
import os
from bs4 import BeautifulSoup
import zipfile

'''
Author: Arnab Mukherjee
version: 3.0
'''


class WordDataExtract:
    '''Class takes care of all data extraction'''

    def __init__(self, js_data):
        self.js_data = js_data  # json data of config

    def list_doc_files(self):
        '''Return the list of doc files in repo'''
        doc_type = "docx"  # type of file to search
        file_list = [i.split("\\")[-1] for i in glob(self.js_data.get("data_path")+"\\*"+doc_type)]  # generate list of
        # files
        return file_list

    # Following section commented for version 2 of code, with deals with xml approach
    # @staticmethod
    # def print_para(paragraphs,start,stop):
    #     '''Prints the content'''
    #     # print(start,stop)
    #     for j in range(start, stop):
    #         print(paragraphs[j].text)
    #     return None
    #
    # def dump_para(self,name,paragraphs,start,stop):
    #     '''Dumps the content'''
    #     # print(name,start,stop)
    #     name = name.split(".")[0] + ".txt"
    #     with open(os.path.join(self.js_data.get("req_out"), name), "w+") as f:
    #         for j in range(start, stop):
    #             f.write(paragraphs[j].text)
    #     print("data extracted")
    #     return None
    #
    # def extract_text(self):
    #     '''Extraction engine function'''
    #     doc_object_list = [(i, Document(os.path.join(self.js_data.get("data_path"), i))) for i in self.list_doc_files()]
    #     for name, i in doc_object_list:
    #         print("File used:", name)
    #         start, stop = 0, 0
    #         for para_index, para in enumerate(i.paragraphs):
    #             if para.style.name.startswith(self.js_data.get("style_search")):
    #                 if para.text in self.js_data.get("extract_section"):
    #                     start = para_index + 1
    #                 else:
    #                     if start > stop:
    #                         stop = para_index
    #                         # self.print_para(i.paragraphs,start,stop)
    #                         self.dump_para(name,i.paragraphs,start,stop)  # dumping the data
    #                         start, stop = 0, 0
    #

    # commented for version 3 of code
    # def xml_approach(self):
    #     '''Extracting text from xml format'''
    #     for i in self.list_doc_files():  # iterating through list of docx files.
    #         print("Using file:",i)
    #         file_path = os.path.join(self.js_data.get("data_path"), i)  # generating file path for each doc file
    #         documents = zipfile.ZipFile(file_path)  # working in zip mode
    #         xml_path = "word/document.xml"  # master xml containing all data
    #         raw_xml_content = documents.read(xml_path)  # reading the xml document
    #         bs_xml_content = BeautifulSoup(raw_xml_content,'xml')  # using bs for parsing
    #         with open(os.path.join(self.js_data.get("req_out"), i.split(".")[0] + ".txt"), "w") as f:
    #             f.write(bs_xml_content.getText())  # saving the data in text file

    def extract_text_section(self):
        '''Extract data from section'''
        text_data = []
        start = False
        for i in self.bs_xml_content.find_all("w:p"):
            if not start:
                for j in i.find_all("w:pStyle"):
                    if j.get("w:val") == self.js_data.get("heading_style") and \
                            i.text == self.js_data.get("heading_name"):
                        start = True

            else:
                if len(i.find_all("w:pStyle")) > 0:
                    for j in i.find_all("w:pStyle"):
                        if j.get("w:val") == self.js_data.get("heading_style"):
                            return text_data
                        else:
                            text_data.append(i.text)
                else:
                    text_data.append(i.text)
        return text_data

    def file_tracer(self):
        '''Extracting text from xml format'''
        for i in self.list_doc_files():  # iterating through list of docx files.
            print("Using file:",i)
            file_path = os.path.join(self.js_data.get("data_path"), i)  # generating file path for each doc file
            documents = zipfile.ZipFile(file_path)  # working in zip mode
            xml_path = "word/document.xml"  # master xml containing all data
            raw_xml_content = documents.read(xml_path)  # reading the xml document
            self.bs_xml_content = BeautifulSoup(raw_xml_content,'xml')  # using bs for parsing
            with open(os.path.join(self.js_data.get("req_out"), i.split(".")[0] + ".txt"), "w") as f:
                for line in self.extract_text_section():
                    f.write(line)
                    f.write("\n")  # saving the data in text file


def main():
    config = "req_reco\\configurations\\data_config.json"  # path of config file
    json_data = None
    with open(config) as f:
        json_data = load(f)
    # WordDataExtract(json_data).extract_text()  # calling the master extract engine
    # WordDataExtract(json_data).xml_approach()  # calling the master extract engine # commented for version 3
    WordDataExtract(json_data).file_tracer()  # calling the master extract engine


if __name__ == "__main__":
    main()
